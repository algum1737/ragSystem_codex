import logging
import re
import zipfile
from pathlib import Path

import docx

from .base_parser import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    @property
    def supported_extensions(self) -> list[str]:
        return [".docx"]

    def parse(self, file_path: Path) -> ParsedDocument:
        try:
            doc = docx.Document(file_path)
        except (zipfile.BadZipFile, Exception) as e:
            raise RuntimeError(f"DOCX 파일 열기 실패: {file_path}") from e

        lines = [p.text or '' for p in doc.paragraphs]
        raw_content = "\n".join(lines)
        content = re.sub(r'\s+', ' ', raw_content).strip()

        if not content:
            logger.warning("빈 콘텐츠 추출됨: %s", file_path)

        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "file_name": file_path.name,
            "file_size_bytes": file_path.stat().st_size,
            "parser": "python-docx",
        }

        source_path = str(file_path.resolve())
        return ParsedDocument(content=content, source_path=source_path, metadata=metadata)
